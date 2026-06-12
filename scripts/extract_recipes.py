"""
Extract recipes from .docx and .pdf files into structured JSON.
Reads Julie_s Recipes Christmas 2025/ and outputs to recipes/.
"""

import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).parent.parent
SOURCE_DIR = ROOT / "Julie_s Recipes Christmas 2025"
OUTPUT_DIR = ROOT / "recipes"
INTRO_FILE = ROOT / "Julie_s Recipes Christmas 2025" / "Cookbook Introduction.docx"

# Explicit section header keywords (used when headings exist)
_SECTION_HEADER = {
    "ingredients": re.compile(r"^ingredient", re.IGNORECASE),
    "instructions": re.compile(
        r"^(direction|instruction|method|preparation|steps|how to make)", re.IGNORECASE
    ),
    "notes": re.compile(r"^(note|tip|variation|serving|storage|make ahead)", re.IGNORECASE),
}

# Content-based classifiers (when no headers are present)
_QUANTITY_START = re.compile(
    r"^(\d[\d/.\-]*\s|[½¼¾⅓⅔⅛⅜⅝⅞]\s|a |an |one |two |three |four |five |six |pinch|dash|handful)",
    re.IGNORECASE,
)
_NOTE_START = re.compile(
    r"^(note[ :—]|tip[ :—]|variation|when |serving suggestion|make ahead|can be made|store |this recipe|originally|i (first|love|always|used|made)|my )",
    re.IGNORECASE,
)
_INSTRUCTION_VERB = re.compile(
    r"^(place|stir|mix|bake|cook|add|combine|heat|bring|pour|whisk|blend|fold|drain|season|serve|remove|let |allow|refrigerate|freeze|preheat|spread|top |brush|cut |chop|dice|slice|mash|cream|beat|dissolve|prepare|in small|in large|in medium|in a |using |for the|meanwhile|transfer|cover|reduce|increase|simmer|boil|roast|grill|fry|saut|toss|coat|arrange|layer|press|form|roll|shape|drop|scoop|set aside|strain|sift|knead|punch|divide|rise|chill|cool|warm|thaw|marinate)",
    re.IGNORECASE,
)


def slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text)
    return text[:80]


def category_slug(name: str) -> str:
    return slugify(name)


def _detect_header(text: str) -> str | None:
    stripped = text.strip().rstrip(":").strip()
    for section, pattern in _SECTION_HEADER.items():
        if pattern.match(stripped):
            return section
    return None


def _classify_line(text: str, current: str) -> str:
    """Classify a line into a section using content heuristics."""
    if _NOTE_START.match(text):
        return "notes"
    if current == "notes":
        return "notes"
    if _QUANTITY_START.match(text):
        return "ingredients"
    # Long sentence starting with instruction verb = likely instructions
    if len(text) > 40 and _INSTRUCTION_VERB.match(text):
        return "instructions"
    # Short line with instruction verb is ambiguous — once in instructions stay there
    if current == "instructions" and _INSTRUCTION_VERB.match(text):
        return "instructions"
    # Once switched to instructions, non-quantity lines stay there
    if current == "instructions" and not _QUANTITY_START.match(text) and len(text) > 20:
        return "instructions"
    return current


def _parse_lines(lines: list[str], title_from_content: bool = False) -> dict:
    """
    Parse a flat list of text lines into recipe sections.
    If title_from_content=True, the first non-empty line is treated as the title.
    """
    title: str | None = None
    current = "ingredients"
    sections: dict[str, list[str]] = {
        "ingredients": [],
        "instructions": [],
        "notes": [],
    }
    header_seen = False

    for raw in lines:
        text = raw.strip()
        if not text:
            continue

        if title is None and title_from_content:
            title = text
            continue

        # Explicit section header?
        detected = _detect_header(text)
        if detected:
            current = detected
            header_seen = True
            continue

        if header_seen:
            # Trust explicit headers — just bucket into current section
            sections[current].append(text)
        else:
            # No headers found yet — use content classifiers
            current = _classify_line(text, current)
            sections[current].append(text)

    return {
        "title": title,
        "ingredients": sections["ingredients"],
        "instructions": sections["instructions"],
        "notes": sections["notes"],
    }


def parse_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(path)
    title: str | None = None
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading" in style and title is None:
            title = text
        else:
            lines.append(text)

    parsed = _parse_lines(lines, title_from_content=(title is None))
    if parsed["title"]:
        title = parsed["title"]
    if title is None:
        title = path.stem

    return {
        "title": title,
        "ingredients": parsed["ingredients"],
        "instructions": parsed["instructions"],
        "notes": parsed["notes"],
    }


def parse_pdf(path: Path) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        lines.extend(text.splitlines())

    parsed = _parse_lines(lines, title_from_content=True)
    title = parsed["title"] or path.stem

    return {
        "title": title,
        "ingredients": parsed["ingredients"],
        "instructions": parsed["instructions"],
        "notes": parsed["notes"],
    }


def extract_intro() -> None:
    if not INTRO_FILE.exists():
        return

    from docx import Document

    doc = Document(INTRO_FILE)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    out = OUTPUT_DIR / "intro.json"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        json.dumps({"paragraphs": paragraphs}, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    print(f"  intro -> {out.relative_to(ROOT)}")


def extract_all() -> int:
    if not SOURCE_DIR.exists():
        print(f"ERROR: source directory not found: {SOURCE_DIR}", file=sys.stderr)
        return 1

    extract_intro()

    total = skipped = errors = 0

    for category_dir in sorted(SOURCE_DIR.iterdir()):
        if not category_dir.is_dir():
            continue

        cat_name = category_dir.name
        cat_slug = category_slug(cat_name)
        out_cat = OUTPUT_DIR / cat_slug
        out_cat.mkdir(parents=True, exist_ok=True)

        recipe_files = list(category_dir.rglob("*.docx")) + list(
            category_dir.rglob("*.pdf")
        )

        for src in sorted(recipe_files):
            # Skip macOS metadata files and Word lock files (~$...)
            if src.name.startswith("._") or src.name.startswith("~$") or src.name == ".DS_Store":
                skipped += 1
                continue

            # Skip the intro file itself
            if src.resolve() == INTRO_FILE.resolve():
                continue

            slug = slugify(src.stem)
            out_file = out_cat / f"{slug}.json"

            try:
                if src.suffix.lower() == ".docx":
                    data = parse_docx(src)
                elif src.suffix.lower() == ".pdf":
                    data = parse_pdf(src)
                else:
                    skipped += 1
                    continue

                data["category"] = cat_name
                data["category_slug"] = cat_slug
                data["slug"] = slug
                data["source_file"] = src.name
                data["photo"] = None

                # Check for matching photo
                photo_path = ROOT / "photos" / f"{slug}.jpg"
                if photo_path.exists():
                    data["photo"] = f"../../photos/{slug}.jpg"

                out_file.write_text(
                    json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8"
                )
                print(f"  {cat_name}/{slug}")
                total += 1

            except Exception as exc:
                print(f"  ERROR {src.name}: {exc}", file=sys.stderr)
                errors += 1

        # .doc files — report as skipped (binary format, needs Word to convert)
        for src in sorted(category_dir.rglob("*.doc")):
            if src.suffix.lower() == ".doc":
                print(
                    f"  SKIP (legacy .doc) {src.name} -- open in Word and Save As .docx to include",
                    file=sys.stderr,
                )
                skipped += 1

    print(f"\nExtracted {total} recipes  |  skipped {skipped}  |  errors {errors}")
    return 0 if errors == 0 else 1


if __name__ == "__main__":
    sys.exit(extract_all())
